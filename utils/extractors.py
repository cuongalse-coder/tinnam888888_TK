"""Module trích xuất nội dung từ các định dạng tài liệu khác nhau.

Hỗ trợ: Excel (.xlsx, .xls), CSV (.csv), PDF (.pdf),
         Hình ảnh (.jpg, .jpeg, .png), Word (.docx).

Mỗi hàm trả về dict chuẩn chứa ``raw_text`` và dữ liệu có cấu trúc,
hoặc ``{'error': str}`` nếu xảy ra lỗi.
"""

from __future__ import annotations

import csv
import io
import re
from pathlib import PurePosixPath
from typing import Any

import pandas as pd
import pdfplumber
import pytesseract
from docx import Document
from PIL import Image


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _safe_seek(file: Any) -> None:
    """Đặt lại con trỏ đọc về đầu nếu có thể."""
    if hasattr(file, "seek"):
        file.seek(0)


def _file_ext(file: Any) -> str:
    """Trả về phần mở rộng file (viết thường, bao gồm dấu chấm)."""
    name: str = getattr(file, "name", "")
    return PurePosixPath(name).suffix.lower()


def _dataframe_to_raw_text(df: pd.DataFrame) -> str:
    """Chuyển DataFrame thành chuỗi text thô (tab-separated)."""
    # Xử lý NaN thành chuỗi rỗng để tránh bị biến thành chữ 'nan'
    df_clean = df.fillna("")
    lines: list[str] = []
    # Header
    cols = [str(c) if str(c) not in ['nan', 'NaN', '<NA>'] else "" for c in df_clean.columns]
    lines.append("\t".join(cols))
    for _, row in df_clean.iterrows():
        vals = [str(v) if str(v) not in ['nan', 'NaN', '<NA>'] else "" for v in row.values]
        lines.append("\t".join(vals))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Excel / CSV
# ---------------------------------------------------------------------------

def extract_excel(file: Any) -> dict[str, Any]:
    """Trích xuất dữ liệu từ file Excel (.xlsx/.xls) hoặc CSV.

    Parameters
    ----------
    file : UploadedFile
        Đối tượng file từ Streamlit (có ``.read()``, ``.name``, ``.type``).

    Returns
    -------
    dict
        ``{'sheets': [...], 'raw_text': str}`` hoặc ``{'error': str}``.
    """
    try:
        _safe_seek(file)
        ext = _file_ext(file)
        sheets: list[dict[str, Any]] = []

        if ext == ".csv":
            try:
                df = pd.read_csv(file)
            except Exception:
                _safe_seek(file)
                try:
                    df = pd.read_csv(file, encoding="utf-8-sig")
                except Exception:
                    _safe_seek(file)
                    df = pd.read_csv(file, encoding="latin-1")
            raw = _dataframe_to_raw_text(df)
            sheets.append({"name": "Sheet1", "dataframe": df, "raw_text": raw})
        else:
            # Excel – đọc tất cả sheet
            try:
                _safe_seek(file)
                engine = "openpyxl" if ext == ".xlsx" else "xlrd"
                xls = pd.ExcelFile(file, engine=engine)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    # Bỏ bớt dòng/cột trống hoàn toàn để tránh loãng
                    df = df.dropna(how="all").dropna(axis=1, how="all")
                    raw = _dataframe_to_raw_text(df)
                    sheets.append({
                        "name": sheet_name,
                        "dataframe": df,
                        "raw_text": raw,
                    })
            except Exception as excel_err:
                # Hầu hết file .xls xuất từ phần mềm hải quan ở Việt Nam thực chất là file HTML chứa thẻ <table>
                try:
                    _safe_seek(file)
                    dfs = pd.read_html(file)
                    if dfs:
                        for idx, df in enumerate(dfs):
                            df = df.dropna(how="all").dropna(axis=1, how="all")
                            raw = _dataframe_to_raw_text(df)
                            sheets.append({
                                "name": f"Bảng {idx+1}",
                                "dataframe": df,
                                "raw_text": raw,
                            })
                    else:
                        raise ValueError("Không tìm thấy bảng dữ liệu HTML")
                except Exception as html_err:
                    # Nếu tất cả thất bại, đọc thô dưới dạng text để trích xuất chuỗi chữ
                    try:
                        _safe_seek(file)
                        content = file.read()
                        raw_text = ""
                        for encoding in ["utf-8", "utf-8-sig", "utf-16", "cp1252", "latin-1"]:
                            try:
                                raw_text = content.decode(encoding)
                                if raw_text.strip():
                                    break
                            except Exception:
                                continue
                        if raw_text.strip():
                            df = pd.DataFrame({"Nội dung trích xuất": [raw_text[:1000]]})
                            sheets.append({
                                "name": "Văn bản gốc",
                                "dataframe": df,
                                "raw_text": raw_text,
                            })
                        else:
                            raise ValueError(f"Không thể giải mã dữ liệu: {excel_err}")
                    except Exception as text_err:
                        return {"error": f"Lỗi định dạng Excel (Chi tiết: {excel_err} | Thử HTML: {html_err} | Thử Text: {text_err})"}

        if not sheets:
            return {"error": "Không thể trích xuất bất kỳ dữ liệu nào từ file Excel này"}

        all_raw = "\n\n".join(s["raw_text"] for s in sheets)
        return {"sheets": sheets, "raw_text": all_raw}

    except Exception as exc:
        return {"error": f"Lỗi khi đọc file Excel/CSV: {exc}"}


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def extract_pdf(file: Any, api_key: str = "", ocr_languages: list[str] | None = None) -> dict[str, Any]:
    """Trích xuất văn bản và bảng từ file PDF.

    Tự động fallback sang OCR (Gemini) nếu trang PDF là ảnh scan
    (text < 50 ký tự).

    Parameters
    ----------
    file : UploadedFile
        Đối tượng file PDF.
    api_key : str
        API key Gemini dùng để OCR ảnh scan.
    ocr_languages : list[str] | None
        Ngôn ngữ cần nhận dạng cho OCR.

    Returns
    -------
    dict
        ``{'pages': [...], 'raw_text': str, 'page_count': int, 'ocr_used': bool}``
        hoặc ``{'error': str}``.
    """
    try:
        _safe_seek(file)
        pages_data: list[dict[str, Any]] = []
        ocr_used = False

        with pdfplumber.open(file) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""

                # Fallback OCR nếu text quá ngắn (ảnh scan)
                if len(text.strip()) < 50:
                    try:
                        from utils.ocr import ocr_pdf_page
                        ocr_text = ocr_pdf_page(page, api_key, ocr_languages=ocr_languages)
                        if len(ocr_text.strip()) > len(text.strip()):
                            text = ocr_text
                            ocr_used = True
                    except Exception:
                        pass  # OCR không khả dụng → giữ text gốc

                # Extract tables from this page
                page_tables = []
                try:
                    tables = page.extract_tables()
                    if tables:
                        for tbl in tables:
                            if tbl:
                                page_tables.append(tbl)
                                # Convert table to text and append
                                for row in tbl:
                                    if row:
                                        cleaned_row = [str(cell).strip() if cell else '' for cell in row]
                                        text += '\n' + '\t'.join(cleaned_row)
                except Exception:
                    pass  # Table extraction not critical

                pages_data.append({
                    "page_num": idx,
                    "text": text,
                    "tables": page_tables,
                })

        raw_text = "\n\n".join(p["text"] for p in pages_data)
        return {
            "pages": pages_data,
            "raw_text": raw_text,
            "page_count": len(pages_data),
            "ocr_used": ocr_used,
        }

    except Exception as exc:
        return {"error": f"Lỗi khi đọc file PDF: {exc}"}


# ---------------------------------------------------------------------------
# Image (OCR)
# ---------------------------------------------------------------------------

def extract_image(
    file: Any,
    languages: list[str] | None = None,
) -> dict[str, Any]:
    """Trích xuất văn bản từ hình ảnh bằng Tesseract OCR.

    Parameters
    ----------
    file : UploadedFile
        Đối tượng file hình ảnh (.jpg, .jpeg, .png).
    languages : list[str] | None
        Danh sách mã ngôn ngữ Tesseract. Mặc định ``['vie', 'eng']``.

    Returns
    -------
    dict
        ``{'raw_text': str, 'confidence': float, 'details': pd.DataFrame}``
        hoặc ``{'error': str}``.
    """
    if languages is None:
        languages = ["vie", "eng"]
    lang_str = "+".join(languages)

    try:
        _safe_seek(file)
        image = Image.open(file)

        raw_text: str = pytesseract.image_to_string(image, lang=lang_str)

        # Chi tiết OCR (bao gồm confidence)
        details_df: pd.DataFrame = pytesseract.image_to_data(
            image, lang=lang_str, output_type=pytesseract.Output.DATAFRAME,
        )

        # Tính confidence trung bình (bỏ qua giá trị -1 = vùng không nhận dạng)
        valid_conf = details_df.loc[details_df["conf"] >= 0, "conf"]
        avg_confidence = float(valid_conf.mean()) if not valid_conf.empty else 0.0

        return {
            "raw_text": raw_text.strip(),
            "confidence": round(avg_confidence, 2),
            "details": details_df,
        }

    except pytesseract.TesseractNotFoundError:
        return {
            "error": (
                "Không tìm thấy Tesseract OCR. Vui lòng cài đặt Tesseract và "
                "đảm bảo đường dẫn được thêm vào biến môi trường PATH.\n"
                "Hướng dẫn: https://github.com/tesseract-ocr/tesseract#installing-tesseract"
            ),
        }
    except Exception as exc:
        return {"error": f"Lỗi khi OCR hình ảnh: {exc}"}


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def extract_word(file: Any) -> dict[str, Any]:
    """Trích xuất văn bản và bảng từ file Word (.docx).

    Parameters
    ----------
    file : UploadedFile
        Đối tượng file Word.

    Returns
    -------
    dict
        ``{'raw_text': str, 'paragraphs': [...], 'tables': [...]}``
        hoặc ``{'error': str}``.
    """
    try:
        _safe_seek(file)
        doc = Document(io.BytesIO(file.read()))

        # Đoạn văn
        paragraphs: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

        # Bảng
        tables: list[dict[str, Any]] = []
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text.strip() for cell in row.cells])

            if rows_data:
                headers = rows_data[0]
                data_rows = rows_data[1:] if len(rows_data) > 1 else []
                tables.append({"headers": headers, "rows": data_rows})

        raw_text = "\n".join(paragraphs)
        if tables:
            for tbl in tables:
                raw_text += "\n" + "\t".join(tbl["headers"])
                for r in tbl["rows"]:
                    raw_text += "\n" + "\t".join(r)

        return {
            "raw_text": raw_text,
            "paragraphs": paragraphs,
            "tables": tables,
        }

    except Exception as exc:
        return {"error": f"Lỗi khi đọc file Word: {exc}"}


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

_EXT_MAP: dict[str, str] = {
    ".xlsx": "excel",
    ".xls": "excel",
    ".csv": "excel",
    ".pdf": "pdf",
    ".jpg": "image",
    ".jpeg": "image",
    ".png": "image",
    ".docx": "word",
}


def extract_file(
    file: Any,
    ocr_languages: list[str] | None = None,
    api_key: str = "",
) -> dict[str, Any]:
    """Tự động nhận diện loại file và trích xuất nội dung.

    Parameters
    ----------
    file : UploadedFile
        Đối tượng file từ Streamlit.
    ocr_languages : list[str] | None
        Ngôn ngữ OCR (chỉ dùng cho ảnh). Mặc định ``['vie', 'eng']``.

    Returns
    -------
    dict
        ``{'file_type': str, 'data': dict, 'raw_text': str, 'error': str|None}``.
    """
    ext = _file_ext(file)
    file_type = _EXT_MAP.get(ext)

    if file_type is None:
        supported = ", ".join(sorted(_EXT_MAP.keys()))
        return {
            "file_type": "unknown",
            "data": {},
            "raw_text": "",
            "error": (
                f"Định dạng file '{ext}' không được hỗ trợ. "
                f"Các định dạng hỗ trợ: {supported}"
            ),
        }

    extractors = {
        "excel": extract_excel,
        "pdf": lambda f: extract_pdf(f, api_key, ocr_languages),
        "image": lambda f: extract_image(f, languages=ocr_languages),
        "word": extract_word,
    }

    try:
        data = extractors[file_type](file)

        if "error" in data:
            return {
                "file_type": file_type,
                "data": {},
                "raw_text": "",
                "error": data["error"],
            }
            
        raw_text = data.get("raw_text", "")
        # Lọc sạch các từ khóa rác nan, NaN, NaT do lỗi parser của Pandas sinh ra
        if raw_text:
            raw_text = re.sub(r'\b(nan|NaN|NaT|<NA>)\b', '', raw_text, flags=re.IGNORECASE)
            raw_text = re.sub(r'[^\S\n]+', ' ', raw_text)  # collapse horizontal whitespace only
            raw_text = re.sub(r'\n{3,}', '\n\n', raw_text)  # max 2 consecutive newlines
            raw_text = raw_text.strip()

        return {
            "file_type": file_type,
            "data": data,
            "raw_text": raw_text,
            "error": None,
        }

    except Exception as exc:
        return {
            "file_type": file_type,
            "data": {},
            "raw_text": "",
            "error": f"Lỗi không xác định khi xử lý file: {exc}",
        }
